# utils.py
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from datetime import datetime


def read_excel_and_return_data(excel_filepath):
    """Читает данные из Excel и возвращает их в виде словаря."""
    try:
        workbook = load_workbook(excel_filepath, data_only=True)
        sheet = workbook.active

        data = []
        for row in sheet.iter_rows(min_row=1):
            row_data = []
            for cell in row:
                value = cell.value
                if isinstance(value, datetime):
                    value = value.strftime('%Y-%m-%d')
                row_data.append(value)
            data.append(row_data)

        headers = data[0]
        json_data = []
        for row in data[1:]:
            item = dict(zip(headers, row))
            json_data.append(item)
        return json_data

    except FileNotFoundError:
        print(f"Файл {excel_filepath} не найден.")
        return None  # Возвращаем None в случае ошибки
    except Exception as e:
        print(f"Произошла ошибка: {e}")
        return None


def create_table(doc, headers, data, mechanical_locks_item):
    """
    Создает таблицу в Word документе на основе переданных данных.
    
    Args:
        doc (Document): Объект Word документа.
        headers (dict): Словарь с заголовками таблицы.
        data (list): Список данных для таблицы.
        mechanical_locks_item (dict): Данные о механических блокировках.
    """
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    table.layout = 1

    table.columns[0].width = Cm(1.3)
    table.columns[1].width = Cm(12)
    table.columns[2].width = Cm(1.3)
    table.columns[3].width = Cm(1.3)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx == 0:
                cell.width = Cm(1.3)
            elif idx == 1:
                cell.width = Cm(12)
            elif idx == 2:
                cell.width = Cm(1.3)
            elif idx == 3:
                cell.width = Cm(1.3)
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = headers.get("hdr_name", "Наименование")
    hdr_cells[2].text = headers.get("hdr_unit", "Ед. изм.")
    hdr_cells[3].text = headers.get("hdr_quantity", "Кол-во")

    # Форматирование заголовков таблицы
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True 
     
    # Добавление данных в таблицу
    row_number = 1
    section_number = 1
    stop_adding_data = False 
    
    for item in data:
        if item.get("Скрыть строку, символы /*") == "/*":
            continue
        elif item.get("Скрыть строку, символы /*") is None:
        
            if stop_adding_data:
                continue
            
            if item.get("Структура") is not None:
                if section_number > 5:
                    stop_adding_data = True
                    continue
                # Добавляем новую строку для заголовка секции
                row = table.add_row()
                cells = row.cells
                
                # Нумерация секции (ячейка 0)
                cells[0].text = str(section_number)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
                for paragraph in cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True  # Жирный шрифт
                
                # Заголовок секции (ячейка 1)
                p = cells[1].paragraphs[0]
                run = p.add_run(item.get("Структура"))
                run.bold = True  # Жирный шрифт
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
                
                # Пустые ячейки (ячейки 2 и 3)
                cells[2].text = ""
                cells[3].text = ""
                
                # Увеличиваем номер секции и сбрасываем номер строки
                section_number += 1
                row_number = 1
    
            elif item.get("Опция") is not None and item.get("Примечание 1") is not None and item.get("Примечание 4") is not None:
                if section_number <= 5:
                    if section_number == 5 and row_number > 5:
                        stop_adding_data = True
                        continue
                    row = table.add_row()
                    cells = row.cells
                    cells[0].text = f"{section_number - 1}.{row_number}"
                    cells[1].text = item.get("Опция")
                    cells[2].text = item.get("Примечание 4")
                    cells[3].text = str(item.get("Примечание 1"))
                    
                    # Выравнивание по центру для ячеек 0, 2 и 3
                    for idx, cell in enumerate(cells):
                        if idx in [0, 2, 3]:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
                    row_number += 1
    

    # Добавляем механические блокировки после обработки до секции 4.5
    if mechanical_locks_item:
        row = table.add_row()
        cells = row.cells
        cells[0].text = f"4.6"
        cells[1].text = mechanical_locks_item.get("Опция")
        cells[2].text = mechanical_locks_item.get("Примечание 4")
        cells[3].text = str(mechanical_locks_item.get("Примечание 1"))
        
        for idx, cell in enumerate(cells):
            if idx in [0, 2, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table